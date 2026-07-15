from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import shutil
import tempfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("deliverables/Using Live Session - Formal Guide.docx")


COLORS = {
    "ink": "173044",
    "muted": "52645D",
    "green": "DFF5DF",
    "green_dark": "2F6F4E",
    "brown": "D8C3A5",
    "brown_dark": "6F4E32",
    "cream": "FFF9EF",
    "white": "FFFFFF",
    "line": "CBD8CE",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="CBD8CE", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def add_text(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def paragraph_style(paragraph, before=0, after=6, line=1.15, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_page_background(doc, fill):
    body = doc._element.body
    sect_pr = body.sectPr
    background = OxmlElement("w:background")
    background.set(qn("w:color"), fill)
    doc._element.insert(0, background)
    if sect_pr is not None:
        pg_borders = OxmlElement("w:pgBorders")
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement("w:{}".format(edge))
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "6")
            border.set(qn("w:space"), "16")
            border.set(qn("w:color"), COLORS["line"])
            pg_borders.append(border)
        sect_pr.append(pg_borders)


def add_callout(doc, title, body, fill, accent):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent, "12")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    clear_cell(cell)
    p = cell.paragraphs[0]
    paragraph_style(p, after=4, line=1.15)
    add_text(p, title, bold=True, color=accent, size=11)
    p2 = cell.add_paragraph()
    paragraph_style(p2, after=2, line=1.15)
    add_text(p2, body, color=COLORS["ink"], size=10.5)
    doc.add_paragraph()


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    add_page_background(doc, COLORS["green"])

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(COLORS["ink"])
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(COLORS["green_dark"])
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(COLORS["brown_dark"])

    header = section.header.paragraphs[0]
    header.text = "Pre-Quran Live Session Guide"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string(COLORS["muted"])

    footer = section.footer.paragraphs[0]
    footer.text = "Quraan Academy | Live Session Support"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor.from_string(COLORS["muted"])

    hero = doc.add_table(rows=1, cols=1)
    set_table_width(hero)
    hero.autofit = False
    cell = hero.cell(0, 0)
    set_cell_shading(cell, COLORS["cream"])
    set_cell_border(cell, COLORS["brown"], "14")
    clear_cell(cell)
    p = cell.paragraphs[0]
    paragraph_style(p, after=0, line=1.05)
    add_text(p, "Using Live Session", bold=True, color=COLORS["ink"], size=24)
    p2 = cell.add_paragraph()
    paragraph_style(p2, before=2, after=2, line=1.15)
    add_text(p2, "A formal quick guide for starting, joining, and supporting Pre-Quran live classes.", color=COLORS["green_dark"], size=11.5)

    doc.add_paragraph()

    add_callout(
        doc,
        "Purpose",
        "This guide explains the basic live-session flow: start the class, join audio, use the Current Lesson, and support students with the Virtual Tutor.",
        COLORS["white"],
        COLORS["green_dark"],
    )

    h = doc.add_paragraph(style="Heading 1")
    h.text = "Live Session Workflow"
    paragraph_style(h, before=4, after=8)

    steps = [
        ("Open the Live Sessions page", "From the dashboard, find the scheduled session you want to use."),
        ("Start the class", "If you are the teacher or admin, click Start class. This opens the BigBlueButton live room."),
        ("Wait for the session to load", "The system may also open the Current Lesson in a separate window and the Virtual Tutor in another floating window."),
        ("Join audio", "Inside BigBlueButton, join audio when prompted. Choose microphone if you will speak, or listen only if you are observing."),
        ("Guide the class", "Use the live room to speak with students, use chat, and guide the class."),
    ]
    for title, detail in steps:
        p = doc.add_paragraph(style="List Number")
        paragraph_style(p, after=4, line=1.18)
        add_text(p, title + ": ", bold=True, color=COLORS["ink"], size=10.5)
        add_text(p, detail, color=COLORS["ink"], size=10.5)

    h = doc.add_paragraph(style="Heading 1")
    h.text = "Support Tools"
    paragraph_style(h, before=8, after=8)

    tools = doc.add_table(rows=1, cols=2)
    set_table_width(tools)
    tools.autofit = False
    for col in tools.columns:
        for c in col.cells:
            set_cell_border(c, COLORS["line"], "6")
    hdr = tools.rows[0].cells
    hdr[0].text = "Tool"
    hdr[1].text = "How it is used"
    for c in hdr:
        set_cell_shading(c, COLORS["brown"])
        for p in c.paragraphs:
            paragraph_style(p, after=0, line=1.1)
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    for tool, detail in [
        ("Current Lesson", "Shows the lesson students should follow during the class."),
        ("Virtual Tutor", "Provides extra help. Students can type a question from the live lesson, and the tutor should guide them one step at a time."),
        ("BigBlueButton Room", "The live room where the teacher speaks with students, uses chat, and manages the session."),
    ]:
        row = tools.add_row().cells
        row[0].text = tool
        row[1].text = detail
        for c in row:
            set_cell_shading(c, COLORS["white"])
            set_cell_border(c, COLORS["line"], "6")
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                paragraph_style(p, after=0, line=1.15)

    doc.add_paragraph()
    add_callout(
        doc,
        "During the Class",
        "Keep the BigBlueButton room open. Do not close the Current Lesson or Virtual Tutor unless you no longer need them.",
        COLORS["cream"],
        COLORS["brown_dark"],
    )

    h = doc.add_paragraph(style="Heading 1")
    h.text = "End of Session"
    paragraph_style(h, before=2, after=6)
    p = doc.add_paragraph()
    paragraph_style(p, after=8, line=1.2)
    p.add_run("At the end of the session, the teacher can leave or end the meeting from BigBlueButton.").font.size = Pt(10.5)

    add_callout(
        doc,
        "Quick Reminder",
        "Start the class, join audio, use the Current Lesson, and support students with the Virtual Tutor.",
        COLORS["white"],
        COLORS["green_dark"],
    )

    doc.save(OUT)
    patch_page_background(OUT, COLORS["green"])
    return OUT


def patch_page_background(path, fill):
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        with ZipFile(path, "r") as z:
            z.extractall(tmp)
        settings = tmp / "word" / "settings.xml"
        text = settings.read_text(encoding="utf-8")
        if "<w:displayBackgroundShape/>" not in text:
            text = text.replace("</w:settings>", "<w:displayBackgroundShape/></w:settings>")
        settings.write_text(text, encoding="utf-8")
        rebuilt = path.with_suffix(".tmp.docx")
        with ZipFile(rebuilt, "w", ZIP_DEFLATED) as z:
            for item in tmp.rglob("*"):
                if item.is_file():
                    z.write(item, item.relative_to(tmp).as_posix())
        shutil.move(str(rebuilt), str(path))


if __name__ == "__main__":
    print(build_doc())
