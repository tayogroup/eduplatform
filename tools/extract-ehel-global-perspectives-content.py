"""Extract the Ehel Academy Global Perspectives source packs into a content model.

The packs arrive as `Year <n>-<UTC stamp>-<part>.zip` exports from Google Drive —
the SAME filename shape Science and Computing use. A Downloads folder therefore
holds three subjects under indistinguishable names, so every candidate archive is
classified by what its documents SAY before it is used, and anything that is not
Global Perspectives is refused rather than silently imported.

TWO PACK SHAPES
===============
Years 1-3 are adult-led packs, Years 4-8 are self-study packs:

    Years 1-3   Teacher & Parent Guide   Activity Sheet   Mini-Project & Reflection
    Years 4-8   Lesson   Skills Toolkit   Activities & Discussion   Practice & Reflection

Both are normalised onto four roles — Lesson, Toolkit, Activities, Practice — so
the builder reads one shape. The Years 1-3 Guide is written TO THE GROWN-UP
("Hello, and welcome to your child's very first project"), so its role carries
`voice: "adult"`. The course is self-teaching, so the builder must rewrite that
prose into learner-facing explainers; the checker is the gate on that.

WHAT THIS TOOL DOES NOT DO
==========================
It does not interpret. It records the document structure faithfully — headings,
paragraphs, lists, tables and the emoji callout boxes the packs use — and leaves
every editorial decision to build-ehel-global-perspectives-runtime.js. That split
is what lets the builder be fixed and re-run without re-reading the Word files.

Usage:
    python tools/extract-ehel-global-perspectives-content.py \
        --output outputs/global-perspectives-content/global-perspectives-content-model.json
"""

from __future__ import annotations

import argparse
import io
import json
import re
import sys
import zipfile
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

ZIP_PATTERN = "Year *.zip"

# ── subject classification ───────────────────────────────────────────────────
# Every Global Perspectives document prints "<something> - Global Perspectives"
# in its first two lines, which is the cheapest reliable signal. The two
# rejection patterns are the running heads of the other subjects that ship
# under the same filename shape.
GP_HEAD = re.compile(r"global\s+perspectives", re.I)
COMPUTING_HEAD = re.compile(r"year\s*\d+\s*[-–—]?\s*computing|computing\s*[-–—]\s*unit", re.I)
SCIENCE_HEAD = re.compile(r"year\s*\d+\s*[-–—]?\s*science|science\s*[-–—]\s*unit|scientific\s+terms", re.I)

# ── document roles ───────────────────────────────────────────────────────────
# Order matters: "Activities & Discussion" and "Activity Sheet" both contain
# "activit", and "Practice & Reflection" and "Mini-Project & Reflection" both
# contain "reflection", so the more specific patterns are tested first.
DOC_ROLE_RULES = [
    (re.compile(r"skills?\s*toolkit", re.I), "Toolkit"),
    (re.compile(r"teacher\s*(?:&|and)?\s*parent\s*guide|guide\s+for\s+teachers?", re.I), "Lesson"),
    (re.compile(r"\blesson\b", re.I), "Lesson"),
    (re.compile(r"mini[-\s]*project", re.I), "Practice"),
    (re.compile(r"practice", re.I), "Practice"),
    (re.compile(r"activit(?:y|ies)", re.I), "Activities"),
]
# The Years 1-3 guide talks to the grown-up, not the learner.
ADULT_VOICE = re.compile(r"teacher\s*(?:&|and)?\s*parent\s*guide|guide\s+for\s+teachers?", re.I)

# Labels a document may print as its own name. The unit title is whichever of
# the two title lines is NOT one of these — see infer_unit_title.
DOC_LABEL = re.compile(
    r"^(?:my\s+)?(?:teacher\s*(?:&|and)?\s*parent\s*guide|a\s+guide\s+for\s+teachers?(?:\s+and\s+parents?)?|"
    r"activity\s+sheet|activities\s*(?:&|and)\s*discussion|mini[-\s]*project\s*(?:&|and)\s*reflection|"
    r"lesson|skills?\s+toolkit|practice\s*(?:&|and)\s*reflection)$",
    re.I,
)

# ── callout boxes ────────────────────────────────────────────────────────────
# The packs teach through emoji-led boxes: 🤖 talks to the AI tutor, 💡 holds a
# big idea, 🗣 is a speaking prompt, 🤝/👥 point at the live teacher session,
# 🔍 is a worked look at the skill, 🪞 is a reflection. The box's title is the
# emoji line; its body is the run of bold paragraphs directly beneath it.
CALLOUT_MARKER = re.compile(
    r"^([\U0001F000-\U0001FAFF←-⯿☀-➿️‍]+)\s*(.*)$"
)
# ☐ leads a tick-box option inside an activity, not a callout box.
NOT_A_CALLOUT = re.compile(r"^[☐☑✓★☆]")

# The six Cambridge skills, in the order the packs number them: a "1.2" in a
# Skills Toolkit table means strand 1 (Research), sub-strand 2 (Information
# skills). That is the same order the framework file publishes, so a pack code
# resolves to a real framework code without guessing.
STRAND_SUB_STRANDS = {
    1: ["Rq", "Ri", "Rc", "Rf"],
    2: ["Ap", "Ad", "Ac", "As"],
    3: ["Es", "Ea"],
    4: ["Fc", "Ft", "Fv", "Fl"],
    5: ["Cc", "Ct"],
    6: ["Mi", "Ml"],
}
SKILL_NAMES = {
    1: "Research",
    2: "Analysis",
    3: "Evaluation",
    4: "Reflection",
    5: "Collaboration",
    6: "Communication",
}
SKILL_BY_NAME = {v.lower(): k for k, v in SKILL_NAMES.items()}

# Grade N sits at Stage N; stages 1-6 are Primary 0838, 7-9 Lower Secondary 1129.
def framework_for(stage: int) -> dict:
    if stage <= 6:
        return {"code": "0838", "level": "Cambridge Primary Global Perspectives"}
    return {"code": "1129", "level": "Cambridge Lower Secondary Global Perspectives"}


def clean(value: str) -> str:
    """Normalise whitespace and the non-breaking spaces Word scatters about."""
    return re.sub(r"\s+", " ", (value or "").replace("\xa0", " ")).strip()


def paragraph_is_bold(paragraph: Paragraph) -> bool:
    """True when every run carrying text is bold.

    The packs use bold as a structural signal — headings and callout body lines
    are bold, ordinary prose is not — so this is what separates the two.
    """
    flags = [bool(run.bold) for run in paragraph.runs if (run.text or "").strip()]
    return bool(flags) and all(flags)


def list_kind(paragraph: Paragraph) -> str | None:
    style = (paragraph.style.name or "") if paragraph.style is not None else ""
    if "Number" in style:
        return "numbered"
    if "Bullet" in style or "List Paragraph" in style:
        return "bullet"
    return None


def iter_blocks(document: Document):
    """Paragraphs and tables in document order (python-docx exposes them apart)."""
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def parse_document(data: bytes, doc_name: str, role: str, voice: str) -> dict:
    """Turn one .docx into an ordered list of semantic blocks."""
    document = Document(io.BytesIO(data))
    blocks: list[dict] = []
    pending_list: dict | None = None

    def close_list() -> None:
        nonlocal pending_list
        if pending_list and pending_list["items"]:
            blocks.append(pending_list)
        pending_list = None

    for item in iter_blocks(document):
        if isinstance(item, Table):
            close_list()
            rows = []
            for row in item.rows:
                rows.append([clean("\n".join(p.text for p in cell.paragraphs)) for cell in row.cells])
            rows = [r for r in rows if any(c for c in r)]
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        text = clean(item.text)
        if not text:
            continue

        kind = list_kind(item)
        if kind:
            if not pending_list or pending_list["listType"] != kind:
                close_list()
                pending_list = {"type": "list", "listType": kind, "items": []}
            pending_list["items"].append(text)
            continue
        close_list()

        bold = paragraph_is_bold(item)
        marker_match = CALLOUT_MARKER.match(text)
        if marker_match and not NOT_A_CALLOUT.match(text) and marker_match.group(2).strip():
            blocks.append(
                {
                    "type": "callout",
                    "marker": marker_match.group(1).strip(),
                    "title": clean(marker_match.group(2)),
                    "lines": [],
                }
            )
            continue

        # Lines directly beneath a callout are that callout's body, and the two
        # pack shapes format them differently: the Years 1-3 sheets set the body
        # bold ("🗣 Talk about it" over two bold questions), the Years 4-8 books
        # leave it as plain prose ("🤖 Ask Your AI Tutor" over three sentences).
        # So the run is typed by its FIRST line and continues while the lines
        # match it. Requiring bold dropped every Years 4-8 callout body; taking
        # everything until the next heading swallowed the Years 1-3 headings,
        # which are bold too.
        if blocks and blocks[-1]["type"] == "callout":
            box = blocks[-1]
            if not box["lines"]:
                box["bodyIsBold"] = bold
                box["lines"].append(text)
                continue
            if bold == box.get("bodyIsBold"):
                box["lines"].append(text)
                continue

        blocks.append({"type": "heading" if bold else "paragraph", "text": text})

    close_list()
    return {
        "document": doc_name,
        "role": role,
        "voice": voice,
        "blockCount": len(blocks),
        "blocks": blocks,
    }


def document_role(name: str) -> tuple[str, str]:
    """(role, voice) for a source document filename."""
    stem = re.sub(r"^Year \d+ - Unit \d+ ", "", Path(name).stem)
    voice = "adult" if ADULT_VOICE.search(stem) else "learner"
    for pattern, role in DOC_ROLE_RULES:
        if pattern.search(stem):
            return role, voice
    return "Other", voice


def classify_archive(archive: zipfile.ZipFile) -> str:
    """'GlobalPerspectives', 'Computing', 'Science' or 'Unknown'."""
    votes: defaultdict[str, int] = defaultdict(int)
    for entry in archive.infolist():
        if not entry.filename.lower().endswith(".docx"):
            continue
        try:
            document = Document(io.BytesIO(archive.read(entry)))
        except Exception:
            continue
        head = " ".join(clean(p.text) for p in document.paragraphs[:8])
        if GP_HEAD.search(head):
            votes["GlobalPerspectives"] += 1
        elif COMPUTING_HEAD.search(head):
            votes["Computing"] += 1
        elif SCIENCE_HEAD.search(head):
            votes["Science"] += 1
    if not votes:
        return "Unknown"
    return max(votes.items(), key=lambda kv: kv[1])[0]


def archives_by_year(downloads: Path, pattern: str = ZIP_PATTERN) -> dict[int, Path]:
    """Newest Global Perspectives archive per year.

    Candidates are walked newest-stamp-first and the first that classifies as
    Global Perspectives wins, so a newer Science or Computing export for the
    same year never displaces a Global Perspectives pack.
    """
    candidates: defaultdict[int, list[tuple[str, Path]]] = defaultdict(list)
    for path in downloads.glob(pattern):
        match = re.search(r"Year\s+(\d+)\s*[^-]*-\s*(\d{8}T\d{6}Z)", path.name)
        if not match:
            continue
        candidates[int(match.group(1))].append((match.group(2), path))

    chosen: dict[int, Path] = {}
    for year, entries in sorted(candidates.items()):
        for _, path in sorted(entries, reverse=True):
            with zipfile.ZipFile(path) as archive:
                if classify_archive(archive) == "GlobalPerspectives":
                    chosen[year] = path
                    break
    return chosen


def infer_unit_title(documents: list[dict]) -> str:
    """The unit's topic, as printed at the top of its documents.

    Every document opens with two title lines, but which one carries the topic
    moves between years:

        Year 1   "Year 1 - Unit 1: Teacher & Parent Guide"
                 "What can families teach us? - Global Perspectives"   <- topic
        Year 3   "Year 3 - Unit 2: Rule Makers"                        <- topic
                 "Activity Sheet - Global Perspectives"
        Year 6   "Year 6 - Unit 1: Research"                           <- topic
                 "Research - Global Perspectives"

    So the topic is found by elimination: of the two candidates, take the one
    that is not the document naming itself. Taking line 2 unconditionally is
    what would title three Year 3 units "Activity Sheet".
    """
    votes: defaultdict[str, int] = defaultdict(int)
    for doc in documents:
        heads = [b["text"] for b in doc["blocks"][:4] if b["type"] in ("heading", "paragraph")][:2]
        if not heads:
            continue
        candidates = []
        first = re.sub(r"^Year\s+\d+\s*[-–—]\s*Unit\s+\d+\s*:\s*", "", heads[0]).strip()
        if first and first != heads[0]:
            candidates.append(first)
        if len(heads) > 1:
            second = re.sub(r"\s*[-–—]\s*Global Perspectives\s*$", "", heads[1]).strip()
            if second and second != heads[1]:
                candidates.append(second)
        for candidate in candidates:
            if not DOC_LABEL.match(candidate):
                votes[candidate] += 1
                break
    if not votes:
        return ""
    return max(votes.items(), key=lambda kv: (kv[1], -len(kv[0])))[0]


def extract_objectives(documents: list[dict], stage: int) -> list[dict]:
    """Cambridge objectives the pack itself quotes, resolved to framework codes.

    A Skills Toolkit prints a `Code | What Cambridge says | What it means for
    you` table whose codes are strand.sub-strand ordinals ("1.2" = Research,
    Information skills). That ordering is the framework's own, so the pack code
    resolves to a real framework code — and because the table also quotes the
    objective text, the builder can prove the resolution rather than trust it.
    """
    found: list[dict] = []
    seen: set[str] = set()
    for doc in documents:
        for block in doc["blocks"]:
            if block["type"] != "table" or not block["rows"]:
                continue
            header = [c.lower() for c in block["rows"][0]]
            if not header or not header[0].startswith("code"):
                continue
            # Two wordings of the same table ship across the years:
            #   Years 7-8   Code | What Cambridge says   | What it means for you
            #   Year 5      Code | What it is called     | What it means for you
            # The middle column is the Cambridge objective in the first and the
            # sub-strand's NAME in the second, so it is only recorded as quoted
            # Cambridge text when it really is one. Requiring "cambridge" in the
            # header is what made Year 5's mapping invisible.
            quotes_cambridge = any("cambridge" in c for c in header)
            names_sub_strand = any("called" in c or "name" in c for c in header)
            if not (quotes_cambridge or names_sub_strand):
                continue
            for row in block["rows"][1:]:
                if len(row) < 2:
                    continue
                match = re.match(r"^(\d)\.(\d)$", row[0].strip())
                if not match:
                    continue
                strand, ordinal = int(match.group(1)), int(match.group(2))
                subs = STRAND_SUB_STRANDS.get(strand) or []
                if not 1 <= ordinal <= len(subs):
                    continue
                code = f"{stage}{subs[ordinal - 1]}.01"
                if code in seen:
                    continue
                seen.add(code)
                found.append(
                    {
                        "code": code,
                        "packCode": row[0].strip(),
                        "skill": SKILL_NAMES[strand],
                        "cambridgeText": row[1].strip() if quotes_cambridge else "",
                        "subStrandName": row[1].strip() if names_sub_strand else "",
                        "learnerText": row[2].strip() if len(row) > 2 else "",
                        "source": doc["document"],
                    }
                )
    return found


# The packs were preserved here after they were cleared out of Downloads, which
# broke this tool completely: the built course and the review overrides live in
# git, but nothing could be rebuilt from source. This is the durable copy, in
# the same extracted-.docx shape inputs/ehel-grade*-source already use.
PRESERVED = Path(__file__).resolve().parent.parent / "inputs" / "ehel-global-perspectives-source"


def documents_by_year(downloads: Path) -> tuple[dict[int, list[tuple[str, bytes]]], dict[int, str]]:
    """Every source document per year, from the Drive exports or the preserved tree.

    Downloads is tried first so a fresh export always wins — re-exporting a year
    and re-running is the normal way to fix a short pack. The preserved copy is
    the fallback, so the pipeline still runs once the zips have been tidied away.
    """
    chosen = archives_by_year(downloads)
    if chosen:
        out: dict[int, list[tuple[str, bytes]]] = {}
        for year, archive_path in sorted(chosen.items()):
            with zipfile.ZipFile(archive_path) as archive:
                out[year] = [
                    (entry.filename, archive.read(entry))
                    for entry in sorted(archive.infolist(), key=lambda e: e.filename)
                    if entry.filename.lower().endswith(".docx")
                ]
        return out, {year: path.name for year, path in chosen.items()}

    if PRESERVED.is_dir():
        # The tree records which archive each year came from, so provenance in
        # the built units stays the export's own filename rather than becoming
        # the name of this directory.
        manifest = {}
        manifest_file = PRESERVED / "source-manifest.json"
        if manifest_file.exists():
            manifest = json.loads(manifest_file.read_text(encoding="utf-8")).get("archives", {})
        out = {}
        origins: dict[int, str] = {}
        for year_dir in sorted(PRESERVED.glob("Year *")):
            match = re.search(r"Year\s+(\d+)", year_dir.name)
            if not match:
                continue
            files = sorted(year_dir.rglob("*.docx"))
            if files:
                year = int(match.group(1))
                out[year] = [(str(f.relative_to(year_dir.parent)), f.read_bytes()) for f in files]
                origins[year] = manifest.get(str(year), f"{PRESERVED.name}/{year_dir.name}")
        if out:
            return out, origins
    return {}, {}


def extract(downloads: Path) -> dict:
    by_year, origins = documents_by_year(downloads)
    if not by_year:
        raise SystemExit(
            f"error: no Global Perspectives source found.\n"
            f"       Looked at every '{ZIP_PATTERN}' in {downloads} and classified each by\n"
            f"       its contents, then at {PRESERVED}."
        )

    grades = []
    for year, documents in sorted(by_year.items()):
        units: defaultdict[int, list[dict]] = defaultdict(list)
        for filename, data in documents:
            match = re.search(r"Unit\s+(\d+)", filename)
            if not match:
                continue
            name = Path(filename).name
            role, voice = document_role(name)
            units[int(match.group(1))].append(parse_document(data, name, role, voice))

        stage = year
        unit_models = []
        for unit_no, documents in sorted(units.items()):
            documents.sort(key=lambda d: ["Lesson", "Toolkit", "Activities", "Practice", "Other"].index(d["role"]))
            title = infer_unit_title(documents)
            skill = SKILL_BY_NAME.get(title.lower())
            unit_models.append(
                {
                    "unitNo": unit_no,
                    "unitTitle": title,
                    "skill": SKILL_NAMES[skill] if skill else None,
                    "skillOrder": skill,
                    "packShape": "guided" if any(d["voice"] == "adult" for d in documents) else "self-study",
                    "roles": sorted({d["role"] for d in documents}),
                    "objectives": extract_objectives(documents, stage),
                    "blockCount": sum(d["blockCount"] for d in documents),
                    "documents": documents,
                }
            )

        grades.append(
            {
                "year": year,
                "stage": stage,
                "framework": framework_for(stage),
                "sourceArchive": origins.get(year, ""),
                "unitCount": len(unit_models),
                "units": unit_models,
            }
        )

    return {
        "schemaVersion": "Ehel Global Perspectives Content Model v1.0",
        "subject": "Global Perspectives",
        "generatedAt": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "downloads": str(downloads),
        "gradeCount": len(grades),
        "grades": grades,
    }


def report(model: dict) -> None:
    print(f"Global Perspectives content model - {model['gradeCount']} grade(s)")
    for grade in model["grades"]:
        shapes = sorted({u["packShape"] for u in grade["units"]})
        missing = [u["unitNo"] for u in grade["units"] if not u["unitTitle"]]
        print(
            f"  Year {grade['year']} (Stage {grade['stage']}, {grade['framework']['code']}): "
            f"{grade['unitCount']} unit(s), {sum(u['blockCount'] for u in grade['units'])} blocks, "
            f"{'/'.join(shapes)}  [{grade['sourceArchive']}]"
        )
        for unit in grade["units"]:
            objectives = f", {len(unit['objectives'])} quoted objective(s)" if unit["objectives"] else ""
            print(
                f"      Unit {unit['unitNo']}: {unit['unitTitle'] or '(NO TITLE)'} "
                f"- {'/'.join(unit['roles'])}{objectives}"
            )
        if missing:
            print(f"      WARNING: no title inferred for unit(s) {missing}")


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--downloads", type=Path, default=Path.home() / "Downloads")
    ap.add_argument("--output", type=Path, required=True)
    ap.add_argument("--quiet", action="store_true")
    args = ap.parse_args()

    model = extract(args.downloads)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(model, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    if not args.quiet:
        report(model)
    print(f"wrote {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
