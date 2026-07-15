from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "tmp" / "year2-curriculum" / "Year 2"
DEFAULT_OUTPUT = ROOT / "src" / "prototypes" / "ehel-academy" / "vocabulary" / "grade2-vocabulary.json"

UNIT_VISUALS = {
    1: {"image": "./assets/unit-1-welcome-calendar.png", "alt": "Children greeting one another beside a colourful classroom calendar"},
    2: {"image": "./assets/unit-2-neighbours-jobs.png", "alt": "Children meeting helpful workers in their neighbourhood"},
    3: {"image": "./assets/unit-3-ready-steady-go.png", "alt": "Children moving, exercising and naming parts of the body"},
    4: {"image": "./assets/unit-4-big-sky.png", "alt": "Children observing sunlight, shadows, clouds and the night sky"},
    5: {"image": "./assets/unit-5-measure.png", "alt": "Children measuring shapes and comparing classroom objects"},
    6: {"image": "./assets/unit-6-bugs.png", "alt": "Children discovering insects and garden animals"},
    7: {"image": "./assets/unit-7-world-around-us.png", "alt": "Children planting, recycling and caring for nature"},
    8: {"image": "./assets/unit-8-home.png", "alt": "A warm family home showing rooms, furniture and household jobs"},
    9: {"image": "./assets/unit-9-city.png", "alt": "Children exploring city places, transport and an aquarium"},
}

TYPE_INFO = {
    "noun": ("noun", "Names a person, place, thing or idea."),
    "verb": ("verb", "Shows an action or a state."),
    "adjective": ("adjective", "Describes a noun."),
    "adverb": ("adverb", "Tells how, when or where an action happens."),
    "number": ("number", "Names a quantity or position in an order."),
    "position word": ("position", "Shows where one thing is compared with another."),
    "greeting": ("expression", "A word or phrase used when people meet or leave."),
    "phrase": ("phrase", "A small group of words that works together."),
}

COLORS = ["red", "blue", "green", "yellow", "orange", "purple", "black", "white", "pink", "brown"]
CARDINALS = [
    ("one", 1), ("two", 2), ("three", 3), ("four", 4), ("five", 5), ("six", 6),
    ("seven", 7), ("eight", 8), ("nine", 9), ("ten", 10), ("eleven", 11), ("twelve", 12),
]
ORDINALS = [
    "first", "second", "third", "fourth", "fifth", "sixth", "seventh", "eighth", "ninth", "tenth",
    "eleventh", "twelfth", "thirteenth", "fourteenth", "fifteenth", "sixteenth", "seventeenth",
    "eighteenth", "nineteenth", "twentieth", "twenty-first", "twenty-second", "twenty-third",
    "twenty-fourth", "twenty-fifth", "twenty-sixth", "twenty-seventh", "twenty-eighth",
    "twenty-ninth", "thirtieth",
]
WEATHER = {
    "sunny": "Bright with plenty of sunshine.",
    "cloudy": "Covered with many clouds.",
    "windy": "With a lot of moving air.",
    "rainy": "With rain falling or likely to fall.",
    "dry": "With little or no rain or water.",
}
APPRECIATION = {
    "glad": ("adjective", "Feeling pleased or happy."),
    "happy": ("adjective", "Feeling pleased and joyful."),
    "thankful": ("adjective", "Feeling grateful for something good."),
    "appreciate": ("verb", "To notice and be thankful for something good."),
    "important": ("adjective", "Having great value or meaning."),
}

SPLIT_MEANINGS = {
    "hello": "A greeting we say when we meet someone.",
    "goodbye": "A word we say when we leave someone.",
    "sunrise": "The time when the sun comes up in the morning.",
    "sunset": "The time when the sun goes down in the evening.",
    "long": "Reaching a great distance from one end to the other.",
    "short": "Reaching only a small distance from one end to the other.",
    "high": "Far above the ground or another level.",
    "low": "Near the ground or below another level.",
    "bright": "Giving or reflecting a lot of light.",
    "dark": "Having very little light.",
    "blue": "Having the colour of a clear daytime sky.",
    "grey": "Having a colour between black and white.",
    "orange": "Having the colour of a ripe orange fruit.",
    "big": "Large in size.",
    "small": "Little in size.",
    "tall": "High from bottom to top.",
    "heavy": "Having a lot of weight and being hard to lift.",
    "light": "Having little weight and being easy to lift.",
    "wide": "Measuring a large distance from side to side.",
    "narrow": "Measuring only a small distance from side to side.",
}

SPLIT_CONTEXT_MEANINGS = {
    ("tall / short", "short"): "Not tall; having a small height from bottom to top.",
}


def slug(value: str) -> str:
    text = value.lower().replace("'", "")
    return re.sub(r"[^a-z0-9]+", "-", text).strip("-")


def clean_title(value: str) -> str:
    return re.sub(r"^Vocabulary\s+[^A-Za-z0-9]+\s+Unit\s+\d+:\s*", "", value).strip()


def clean_group(value: str) -> str:
    return re.sub(r"^Group\s+\d+:\s*", "", value).strip()


def normalize_type(source_type: str) -> tuple[str, str]:
    lowered = source_type.lower().strip()
    if "phrasal verb" in lowered or "verb phrase" in lowered:
        return "verb", "Shows an action using a group of words."
    if "noun" in lowered and "verb" in lowered:
        return "verb", "Can name an activity or show an action."
    for source, information in TYPE_INFO.items():
        if source in lowered:
            return information
    return "expression", "A useful word or phrase used to share meaning."


def type_sentences(word: str, source_type: str, example: str, group_title: str) -> list[str]:
    normalized_type, _ = normalize_type(source_type)
    topic = group_title.lower()
    capitalized = word[0].upper() + word[1:]
    if normalized_type == "verb" and word.endswith("ing"):
        extras = [
            f"We are {word} together during our class activity.",
            f"{capitalized} is an action we can show with our bodies.",
            f"The children are {word} while the teacher watches.",
            f"Can you make a sentence about someone {word}?",
        ]
    elif normalized_type == "verb":
        extras = [
            f"I can {word} during our class activity.",
            f"We {word} together when it is our turn.",
            f"Please {word} safely and carefully.",
            f"Can you {word} and explain what you did?",
        ]
    elif normalized_type == "adjective":
        extras = [
            f"The picture looks {word}.",
            f"We found something {word} in the story.",
            f"Can you describe an object as {word}?",
            f"I used {word} to add a clear describing detail.",
        ]
    elif normalized_type == "adverb":
        extras = [
            f"The child completed the action {word}.",
            f"Our teacher asked us to speak {word}.",
            f"Can you show how to move {word}?",
            f"I used {word} to explain how the action happened.",
        ]
    elif normalized_type == "number":
        extras = [
            f"Count aloud until you reach {word}.",
            f"We made a group that shows {word}.",
            f"Can you write {word} correctly?",
            f"I heard {word} in our number game.",
        ]
    elif normalized_type == "position":
        extras = [
            f"Place the pencil {word} the book.",
            f"The picture shows one object {word} another object.",
            f"Can you put your hand {word} the table?",
            f"I used {word} to explain where the object is.",
        ]
    else:
        extras = [
            f"We learned the word {word} in our {topic} lesson.",
            f"I wrote {word} in my vocabulary book.",
            f"Can you find or show {word} in the picture?",
            f"I can use {word} in a clear sentence.",
        ]
    return [example, *extras]


def starter_for(word: str, normalized_type: str) -> str:
    if normalized_type == "verb":
        return f"I {word}"
    if normalized_type == "adjective":
        return f"The picture is {word}"
    if normalized_type == "adverb":
        return f"She moved {word}"
    if normalized_type == "position":
        return f"The object is {word}"
    return f"The {word}"


def split_entry(entry: dict[str, str]) -> list[dict[str, str]]:
    if " / " not in entry["word"]:
        return [entry]
    terms = [part.strip() for part in entry["word"].split("/") if part.strip()]
    meanings = [part.strip(" .") for part in re.split(r"\s*/\s*", entry["meaning"])]
    output = []
    for index, term in enumerate(terms):
        equals_match = re.search(rf"\b{re.escape(term)}\s*=\s*([^.]*)", entry["meaning"], flags=re.IGNORECASE)
        context_meaning = SPLIT_CONTEXT_MEANINGS.get((entry["word"].lower(), term.lower()))
        if context_meaning:
            meaning = context_meaning
        elif equals_match:
            meaning = equals_match.group(1).strip()
        elif term.lower() in SPLIT_MEANINGS:
            meaning = SPLIT_MEANINGS[term.lower()]
        else:
            meaning = meanings[index] if len(meanings) == len(terms) else entry["meaning"]
        output.append({**entry, "word": term, "meaning": meaning.rstrip(".") + "."})
    return output


def make_word(unit_number: int, group_number: int, index: int, entry: dict[str, str], group_title: str) -> dict:
    normalized_type, type_definition = normalize_type(entry["type"])
    word = entry["word"].strip()
    return {
        "id": f"u{unit_number}-g{group_number}-{index + 1}-{slug(word)}",
        "word": word,
        "type": normalized_type,
        "sourceType": entry["type"].strip(),
        "typeDefinition": type_definition,
        "pronunciation": "Listen, then repeat",
        "meaning": entry["meaning"].strip(),
        "example": entry["example"].strip(),
        "sentences": type_sentences(word, entry["type"], entry["example"].strip(), group_title),
        "starter": starter_for(word, normalized_type),
        "tutorPrompt": f"Use the word '{word}' in a sentence, then explain what it means.",
    }


def table_entries(document: Document) -> list[list[dict[str, str]]]:
    groups = []
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows or rows[0][:4] != ["Word", "Type", "What it means", "Example"]:
            continue
        groups.append([
            {"word": row[0], "type": row[1], "meaning": row[2], "example": row[3]}
            for row in rows[1:]
            if len(row) >= 4 and row[0]
        ])
    return groups


def supplemental_groups(unit_number: int) -> list[tuple[str, list[dict[str, str]]]]:
    if unit_number == 1:
        colors = [
            {"word": color, "type": "adjective", "meaning": f"The colour {color}.", "example": f"I found something {color} in the room."}
            for color in COLORS
        ]
        numbers = [
            {"word": word, "type": "number", "meaning": f"The number {value}.", "example": f"I can count {value} objects and say {word}."}
            for word, value in CARDINALS
        ]
        ordinals = [
            {"word": word, "type": "number", "meaning": f"The order word for position {index}.", "example": f"The {word} day is marked on our calendar."}
            for index, word in enumerate(ORDINALS, start=1)
        ]
        return [("Colours", colors), ("Numbers 1 to 12", numbers), ("Order Words for Dates", ordinals)]
    if unit_number == 4:
        entries = [
            {"word": word, "type": "adjective", "meaning": meaning, "example": f"Today the weather is {word}."}
            for word, meaning in WEATHER.items()
        ]
        return [("Weather Words", entries)]
    if unit_number == 7:
        entries = [
            {"word": word, "type": word_type, "meaning": meaning, "example": f"I am {word} for the trees, water and clean air." if word != "appreciate" else "I appreciate the clean air and shady trees."}
            for word, (word_type, meaning) in APPRECIATION.items()
        ]
        return [("Words for Saying Thank You to Nature", entries)]
    return []


def import_curriculum(source_directory: Path) -> dict:
    units = []
    for unit_number in range(1, 10):
        source = source_directory / f"Unit {unit_number}" / f"Unit {unit_number} - Vocabulary.docx"
        if not source.exists():
            raise FileNotFoundError(f"Vocabulary source not found: {source}")
        document = Document(source)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        source_groups = [clean_group(text) for text in paragraphs if text.startswith("Group ")]
        extracted_groups = table_entries(document)
        groups = []
        for group_index, entries in enumerate(extracted_groups):
            group_number = group_index + 1
            group_title = source_groups[group_index]
            split_entries = [item for entry in entries for item in split_entry(entry)]
            groups.append({
                "id": f"u{unit_number}-g{group_number}-{slug(group_title)}",
                "number": group_number,
                "title": group_title,
                "words": [
                    make_word(unit_number, group_number, index, entry, group_title)
                    for index, entry in enumerate(split_entries)
                ],
            })
        for title, entries in supplemental_groups(unit_number):
            group_number = len(groups) + 1
            groups.append({
                "id": f"u{unit_number}-g{group_number}-{slug(title)}",
                "number": group_number,
                "title": title,
                "words": [
                    make_word(unit_number, group_number, index, entry, title)
                    for index, entry in enumerate(entries)
                ],
            })
        unit_title = clean_title(paragraphs[0])
        units.append({
            "id": f"unit-{unit_number}",
            "number": unit_number,
            "title": unit_title,
            "source": source.name,
            "visual": UNIT_VISUALS[unit_number],
            "groups": groups,
            "wordCount": sum(len(group["words"]) for group in groups),
        })
    return {
        "schemaVersion": 1,
        "school": "Ehel Academy",
        "subject": "English",
        "grade": 2,
        "voice": {"provider": "ElevenLabs", "voiceId": "XfNU2rGpBa01ckF309OY", "modelId": "eleven_multilingual_v2"},
        "units": units,
        "unitCount": len(units),
        "totalWords": sum(unit["wordCount"] for unit in units),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Import the Ehel Academy Grade 2 vocabulary curriculum.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    curriculum = import_curriculum(args.source.resolve())
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(curriculum, indent=2, ensure_ascii=True) + "\n", encoding="utf-8")
    print(f"Wrote {args.out}")
    for unit in curriculum["units"]:
        print(f"Unit {unit['number']}: {len(unit['groups'])} groups, {unit['wordCount']} words")
    print(f"Total: {curriculum['totalWords']} word cards")


if __name__ == "__main__":
    main()
