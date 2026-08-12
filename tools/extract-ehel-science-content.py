from __future__ import annotations

import argparse
import hashlib
import io
import json
import re
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


# Google-Drive exports are stamped "Year <n>-<UTC timestamp>-<part>.zip". A new
# export does not replace the old one, so several stamps for the same year can
# sit in Downloads at once. Match them all and keep only the newest per year --
# pinning a single stamp silently rebuilds from stale source when content is
# re-exported.
ZIP_PATTERN = "Year *.zip"

# ── subject classification ───────────────────────────────────────────────────
# Science, Computing and Global Perspectives all export from Google Drive as
# "Year <n>-<stamp>-<part>.zip", so a Downloads folder holds three subjects
# under indistinguishable names and the newest stamp for a year is often not
# ours. Picking by name alone therefore does not merely risk the wrong pack, it
# actively prefers it: at the time this was written Downloads held a Year 4 and
# a Year 5 Global Perspectives export, both stamped later than their Science
# counterparts, so a plain run would have rebuilt Grades 4 and 5 of Science out
# of another subject's content.
#
# Every document names its subject in its opening lines — "Year 4 Science -
# Unit 2 - Quick Revision Guide" against "Year 4 - Unit 3: Skills Toolkit
# Evaluation - Global Perspectives" — which is the cheapest reliable signal.
# The Computing and Global Perspectives extractors already classify this way;
# this brings Science in line with them.
SCIENCE_HEAD = re.compile(r"year\s*\d+\s*[-–—]?\s*science|science\s*[-–—]\s*unit|scientific\s+terms", re.I)
GP_HEAD = re.compile(r"global\s+perspectives", re.I)
COMPUTING_HEAD = re.compile(r"year\s*\d+\s*[-–—]?\s*computing|computing\s*[-–—]\s*unit", re.I)


def classify_archive(archive: zipfile.ZipFile) -> str:
    """'Science', 'GlobalPerspectives', 'Computing' or 'Unknown'.

    Every .docx votes on the subject named in its first eight paragraphs and the
    majority wins, so one oddly-headed document cannot carry an archive.
    """
    votes: Counter[str] = Counter()
    for entry in archive.infolist():
        if not entry.filename.lower().endswith(".docx"):
            continue
        try:
            document = Document(io.BytesIO(archive.read(entry)))
        except Exception:
            continue
        head = " ".join(clean(paragraph.text) for paragraph in document.paragraphs[:8])
        if GP_HEAD.search(head):
            votes["GlobalPerspectives"] += 1
        elif COMPUTING_HEAD.search(head):
            votes["Computing"] += 1
        elif SCIENCE_HEAD.search(head):
            votes["Science"] += 1
    if not votes:
        return "Unknown"
    return votes.most_common(1)[0][0]


def newest_archive_per_year(downloads: Path, pattern: str = ZIP_PATTERN) -> list[Path]:
    """Newest SCIENCE archive per year.

    Candidates are walked newest-stamp-first and the first that classifies as
    Science wins, so a later export of another subject never displaces a Science
    pack. A year whose archives are all another subject is reported rather than
    skipped in silence: the alternative is a grade quietly vanishing from the
    model and, one build later, from the course.
    """
    candidates: defaultdict[int, list[tuple[str, Path]]] = defaultdict(list)
    for path in downloads.glob(pattern):
        match = re.search(r"Year\s+(\d+)\s*[^-]*-\s*(\d{8}T\d{6}Z)", path.name)
        if not match:
            continue
        candidates[int(match.group(1))].append((match.group(2), path))

    chosen: list[Path] = []
    for year, entries in sorted(candidates.items()):
        rejected: list[tuple[Path, str]] = []
        for _, path in sorted(entries, reverse=True):
            with zipfile.ZipFile(path) as archive:
                subject = classify_archive(archive)
            if subject == "Science":
                chosen.append(path)
                for other, other_subject in rejected:
                    print(f"  Year {year}: skipped {other.name} - it is {other_subject}, not Science")
                break
            rejected.append((path, subject))
        else:
            print(f"  Year {year}: NO Science archive found among {len(entries)} candidate(s):")
            for other, other_subject in rejected:
                print(f"      {other.name} is {other_subject}")
    return chosen


def clean(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def slug(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")


def iter_blocks(document: DocumentObject):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, document)


def document_type(name: str) -> str:
    stem = Path(name).stem
    if "Teacher Guide" in stem:
        return "Lesson"
    if "Activity Sheet" in stem:
        return "Activities"
    for value in ("Lesson", "Activities", "Experiments", "Practice", "Reference"):
        if re.search(rf"\b{value}\b", stem, re.IGNORECASE):
            return value
    return "Other"


def looks_like_heading(text: str, style: str) -> bool:
    if style.lower().startswith(("heading", "title", "subtitle")):
        return True
    if len(text) > 120 or text.endswith((".", "?", "!", ";")):
        return False
    return bool(
        re.match(
            r"^(unit\s+\d+|unit\s+overview|lesson|activity|experiment\s+\d+|practice|reference|objectives?|learning|about\s+this\s+unit|"
            r"key\s+(?:words|terms|vocabulary|rules|science\s+words)|vocabulary|warm[- ]?up|introduction|"
            r"worked\s+examples?|guided\s+practice|independent\s+practice|examples?|"
            r"step[- ]by[- ]step|common\s+mistakes?|section\s+[a-e]|section\s+\d+|"
            r"what\s+you\s+will|aim\b|make\s+a\s+hypothesis|hypothesis\b|method\b|recording\s+sheet|"
            r"analysis\s+questions?|diagram\s+in\s+words|the\s+most\s+important|connections?\s+to|"
            r"end[- ]of[- ]unit\s+quiz|going\s+further|safety|part\s+\d+|"
            r"summary|remember|challenge|assessment|answer\s+key|answers?|teacher|resources?|materials?)\b",
            text,
            re.IGNORECASE,
        )
    )


def content_kind(text: str, style: str, section: str, doc_type: str) -> str:
    lower = f"{section} {text}".lower()
    if looks_like_heading(text, style):
        return "Heading"
    if "answer key" in section.lower() or section.lower().startswith("answer"):
        return "Answer guidance"
    if "objective" in lower or "learning outcome" in lower or "success criteria" in lower:
        return "Learning outcome"
    if "worked example" in lower or re.match(r"^(example|e\.g\.)\b", text, re.IGNORECASE):
        return "Worked example"
    if "key term" in lower or "vocabulary" in lower or "glossary" in lower:
        return "Key term"
    if doc_type in {"Practice", "Activities"} or text.endswith("?"):
        return "Task"
    if style.lower().startswith("list"):
        return "List item"
    return "Instructional text"


def parse_document(data: bytes, source_name: str, grade: int, unit: int, doc_type: str) -> dict:
    document = Document(io.BytesIO(data))
    blocks = []
    section = "Document opening"
    sequence = 0
    for block_type, block in iter_blocks(document):
        if block_type == "paragraph":
            text = clean(block.text)
            if not text:
                continue
            style = block.style.name if block.style else "Normal"
            if looks_like_heading(text, style):
                section = text
            sequence += 1
            blocks.append(
                {
                    "sequence": sequence,
                    "block_type": "Paragraph",
                    "section": section,
                    "content_kind": content_kind(text, style, section, doc_type),
                    "style": style,
                    "text": text,
                    "table_row": None,
                    "table_col": None,
                }
            )
        else:
            for row_number, row in enumerate(block.rows, 1):
                for col_number, cell in enumerate(row.cells, 1):
                    text = clean(cell.text)
                    if not text:
                        continue
                    sequence += 1
                    blocks.append(
                        {
                            "sequence": sequence,
                            "block_type": "Table cell",
                            "section": section,
                            "content_kind": content_kind(text, "Table", section, doc_type),
                            "style": "Table",
                            "text": text,
                            "table_row": row_number,
                            "table_col": col_number,
                        }
                    )
    return {
        "grade": grade,
        "unit": unit,
        "document_type": doc_type,
        "source_file": source_name,
        "sha256": hashlib.sha256(data).hexdigest(),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "block_count": len(blocks),
        "blocks": blocks,
    }


def infer_unit_title(documents: list[dict], unit: int) -> str:
    priority = {"Lesson": 0, "Reference": 1, "Activities": 2, "Practice": 3, "Other": 4}
    candidates = []
    for document in sorted(documents, key=lambda item: priority.get(item["document_type"], 9)):
        if document["blocks"]:
            opening = document["blocks"][0]["text"]
            if 1 < len(opening) <= 100 and not re.match(r"^(year|unit|lesson|practice|activities?|reference|teacher)", opening, re.I):
                candidates.append(opening)
        for block in document["blocks"][:20]:
            text = block["text"]
            match = re.match(rf"^Unit\s+{unit}\s*[:\-–—]\s*(.+)$", text, re.IGNORECASE)
            if match:
                candidate = clean(match.group(1))
                if candidate and not re.match(r"^(lesson|practice|activities?|reference|teacher guide)$", candidate, re.I):
                    candidates.append(candidate)
            if block["content_kind"] == "Heading" and not re.match(r"^(year|unit|lesson|practice|activities?|reference)", text, re.I):
                candidates.append(text)
        if candidates:
            break
    return candidates[0] if candidates else f"Unit {unit}"


def extract(downloads: Path) -> dict:
    grades: dict[int, dict] = {}
    for zip_path in newest_archive_per_year(downloads):
        grade_match = re.search(r"Year\s+(\d+)", zip_path.name)
        if not grade_match:
            continue
        grade = int(grade_match.group(1))
        documents = []
        with zipfile.ZipFile(zip_path) as archive:
            for entry in sorted(archive.infolist(), key=lambda item: item.filename):
                if not entry.filename.lower().endswith(".docx"):
                    continue
                unit_match = re.search(r"Unit\s+(\d+)", entry.filename, re.IGNORECASE)
                if not unit_match:
                    continue
                unit = int(unit_match.group(1))
                doc_type = document_type(entry.filename)
                documents.append(parse_document(archive.read(entry), entry.filename, grade, unit, doc_type))
        by_unit = defaultdict(list)
        for document in documents:
            by_unit[document["unit"]].append(document)
        units = []
        for unit_number in sorted(by_unit):
            unit_docs = by_unit[unit_number]
            counts = Counter(item["document_type"] for item in unit_docs)
            units.append(
                {
                    "grade": grade,
                    "unit": unit_number,
                    "unit_id": f"sci-g{grade:02d}-u{unit_number:02d}",
                    "term": ((unit_number - 1) // 6) + 1,
                    "title": infer_unit_title(unit_docs, unit_number),
                    "lesson_documents": counts["Lesson"],
                    "activity_documents": counts["Activities"],
                    "practice_documents": counts["Practice"],
                    "reference_documents": counts["Reference"],
                    "other_documents": counts["Other"],
                    "source_document_count": len(unit_docs),
                    "source_block_count": sum(item["block_count"] for item in unit_docs),
                    "review_status": "Imported - curriculum review required",
                }
            )
        grades[grade] = {
            "metadata": {
                "schema_version": "Ehel Science Content Package v1.0",
                "grade": grade,
                "grade_id": f"g{grade:02d}",
                "subject": "Science",
                "source_archive": zip_path.name,
                "unit_count": len(units),
                "source_document_count": len(documents),
            },
            "units": units,
            "documents": documents,
        }
    return {"grades": grades}


def sample(model: dict) -> None:
    for grade, package in model["grades"].items():
        print(f"\n### Grade {grade}")
        seen = set()
        for document in package["documents"]:
            doc_type = document["document_type"]
            if doc_type in seen:
                continue
            seen.add(doc_type)
            print(f"\n[{doc_type}] {document['source_file']}")
            for block in document["blocks"][:18]:
                print(f"{block['sequence']:02d} {block['style']} | {block['section']} | {block['text'][:180]}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--downloads", type=Path, default=Path.home() / "Downloads")
    parser.add_argument("--output", type=Path)
    parser.add_argument("--sample", action="store_true")
    args = parser.parse_args()
    model = extract(args.downloads)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"Wrote {args.output}")
    if args.sample:
        sample(model)
    print(json.dumps({grade: package["metadata"] for grade, package in model["grades"].items()}, indent=2))


if __name__ == "__main__":
    main()
