"""Extract the Ehel Academy Computing source packs into a content model.

The Computing packs arrive as Google-Drive exports named exactly like the
Science packs ("Year <n>-<UTC stamp>-<part>.zip"), so a Downloads folder holds
both subjects under indistinguishable filenames. Picking archives by name alone
would silently rebuild Computing from the Science books, so every archive is
classified by what its documents say before it is used.

Usage:
  python tools/extract-ehel-computing-content.py --output outputs/computing-content/computing-content-model.json
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import re
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


ZIP_PATTERN = "Year *.zip"

# A Computing book names its subject in the running head of every document
# ("Year 3 Computing - Unit 5 - Teacher Guide"). Reference and Practice
# booklets lead with their own title instead, so the classifier also accepts a
# clear run of computing vocabulary.
COMPUTING_HEAD = re.compile(r"year\s*\d+\s*[-–—]?\s*computing|computing\s*[-–—]\s*unit", re.I)
SCIENCE_HEAD = re.compile(r"year\s*\d+\s*[-–—]?\s*science|science\s*[-–—]\s*unit|scientific\s+terms", re.I)
COMPUTING_WORDS = re.compile(
    r"\b(scratch|scratchjr|micro:?bit|makecode|algorithm|debug|sprite|pseudocode|"
    r"input device|selection|iteration|variable|spreadsheet|database|network|"
    r"e-safety|program(?:me)?s?\b|app\b|keyboard|binary)\b",
    re.I,
)

# Two document sets ship across the years. Years 1-4 are adult-led packs
# (Teacher Guide + Activity Sheet + Mini-Project); Years 5-7 are self-study
# packs (Lesson + Activities & Projects + Practice & Quiz + Reference). Both
# are normalised onto the same four roles so the builder has one shape to read.
DOC_TYPE_RULES = [
    (re.compile(r"teacher\s*(?:&|and)?\s*parent\s*guide|teacher\s*guide", re.I), "Lesson"),
    (re.compile(r"\blesson\b", re.I), "Lesson"),
    (re.compile(r"activit(?:y|ies)\s*(?:sheet|&|and)?\s*(?:projects?)?", re.I), "Activities"),
    (re.compile(r"mini[- ]project", re.I), "Practice"),
    (re.compile(r"practice", re.I), "Practice"),
    (re.compile(r"reference", re.I), "Reference"),
]

# Headings used by the Computing books. Years 1-4 teach in "Session N" blocks;
# Years 5-7 teach in decimal subunits ("1.3 Selection").
HEADING_PATTERN = re.compile(
    r"^(unit\s+\d+|unit\s+overview|about\s+this\s+unit|how\s+the\s+\w+\s+documents?|"
    r"learning\s+objectives?|objectives?|key\s+(?:words|terms|vocabulary)|glossary|"
    r"materials?\s+needed|what\s+you\s+need|free\s+software|setup\s+guides?|"
    r"lesson\s+flow|session\s+\d+|part\s+[a-e]\b|part\s+\d+|task\s+\d+|step\s+\d+|"
    r"section\s+[a-e]\b|section\s+\d+|activity\s+\d+|project\b|mini[- ]project|"
    r"end[- ]of[- ]unit|challenge\b|extension\b|"
    r"common\s+(?:misconceptions?|errors?|mistakes?|bugs?)|debugging\b|"
    r"assessment|observation\s+checklist|success\s+criteria|home\s+connection|"
    r"answer\s+keys?|answers?\b|practice\s+(?:questions?|answers?)|quiz\b|"
    r"your\s+scenario|scenario\b|cheat[- ]sheet|reference\b|"
    r"e[- ]?safety|staying\s+safe|did\s+i\s+do\s+it|what\s+can\s+you\s+do|"
    r"how\s+to\s+use\s+this|teacher\s+answer\s+key|going\s+further|checklist)\b",
    re.I,
)

# "1.3 Selection" / "2.4 Sorting data" — the Years 5-7 subunit heading.
SUBUNIT_PATTERN = re.compile(r"^(\d{1,2}\.\d{1,2})\s+(\S.{2,90})$")


def clean(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def iter_blocks(document: DocumentObject):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, document)


def document_type(name: str) -> str:
    stem = Path(name).stem
    for pattern, value in DOC_TYPE_RULES:
        if pattern.search(stem):
            return value
    return "Other"


def looks_like_heading(text: str, style: str) -> bool:
    if style.lower().startswith(("heading", "title", "subtitle")):
        return True
    if len(text) > 120 or text.endswith((".", "?", "!", ";", ":")):
        return False
    if SUBUNIT_PATTERN.match(text):
        return True
    return bool(HEADING_PATTERN.match(text))


def content_kind(text: str, style: str, section: str, doc_type: str) -> str:
    lower = f"{section} {text}".lower()
    if looks_like_heading(text, style):
        return "Heading"
    if re.search(r"answer key|teacher answer", section, re.I):
        return "Answer guidance"
    if "objective" in lower or "success criteria" in lower or "what can you do" in lower:
        return "Learning outcome"
    if re.search(r"key (words|terms)|glossary|vocabulary", lower):
        return "Key term"
    if re.match(r"^(what to teach|how to explain)", text, re.I):
        return "Teaching note"
    if re.search(r"worked example|example:", lower) or re.match(r"^(example|e\.g\.)\b", text, re.I):
        return "Worked example"
    if doc_type in {"Practice", "Activities"} or text.endswith("?"):
        return "Task"
    if style.lower().startswith("list"):
        return "List item"
    return "Instructional text"


def parse_document(data: bytes, source_name: str, grade: int, unit: int, doc_type: str) -> dict:
    document = Document(io.BytesIO(data))
    blocks = []
    section = "Document opening"
    sequence = 0
    for block_type, block in iter_blocks(document):
        if block_type == "paragraph":
            text = clean(block.text)
            if not text:
                continue
            style = block.style.name if block.style else "Normal"
            if looks_like_heading(text, style):
                section = text
            sequence += 1
            blocks.append({
                "sequence": sequence,
                "block_type": "Paragraph",
                "section": section,
                "content_kind": content_kind(text, style, section, doc_type),
                "style": style,
                "text": text,
                "table_row": None,
                "table_col": None,
            })
        else:
            for row_number, row in enumerate(block.rows, 1):
                for col_number, cell in enumerate(row.cells, 1):
                    text = clean(cell.text)
                    if not text:
                        continue
                    sequence += 1
                    blocks.append({
                        "sequence": sequence,
                        "block_type": "Table cell",
                        "section": section,
                        "content_kind": content_kind(text, "Table", section, doc_type),
                        "style": "Table",
                        "text": text,
                        "table_row": row_number,
                        "table_col": col_number,
                    })
    return {
        "grade": grade,
        "unit": unit,
        "document_type": doc_type,
        "source_file": source_name,
        "sha256": hashlib.sha256(data).hexdigest(),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "block_count": len(blocks),
        "blocks": blocks,
    }


def classify_archive(archive: zipfile.ZipFile) -> str:
    """Return 'Computing', 'Science' or 'Unknown' for an archive of .docx files."""
    computing = science = 0
    for entry in archive.infolist():
        if not entry.filename.lower().endswith(".docx"):
            continue
        try:
            document = Document(io.BytesIO(archive.read(entry)))
        except Exception:
            continue
        head = " ".join(clean(p.text) for p in document.paragraphs[:8])
        if SCIENCE_HEAD.search(head):
            science += 1
        elif COMPUTING_HEAD.search(head):
            computing += 1
        else:
            body = " ".join(clean(p.text) for p in document.paragraphs[:120])
            if len(COMPUTING_WORDS.findall(body)) >= 4:
                computing += 1
    if computing and not science:
        return "Computing"
    if science and not computing:
        return "Science"
    return "Unknown" if not (computing or science) else ("Computing" if computing > science else "Science")


def computing_archives(downloads: Path, pattern: str = ZIP_PATTERN) -> dict[int, Path]:
    """Newest Computing archive per year.

    Candidates are walked newest-stamp-first and the first one that classifies
    as Computing wins, so a newer Science export for the same year never
    displaces a Computing pack.
    """
    by_year: dict[int, list[tuple[str, Path]]] = defaultdict(list)
    for candidate in downloads.glob(pattern):
        match = re.search(r"Year\s+(\d+)\s*[^-]*-\s*(\d{8}T\d{6}Z)", candidate.name)
        if not match:
            continue
        by_year[int(match.group(1))].append((match.group(2), candidate))

    chosen: dict[int, Path] = {}
    for year, candidates in sorted(by_year.items()):
        for _, path in sorted(candidates, reverse=True):
            with zipfile.ZipFile(path) as archive:
                if classify_archive(archive) == "Computing":
                    chosen[year] = path
                    break
    return chosen


def infer_unit_title(documents: list[dict], unit: int) -> str:
    """The unit's own title, as printed at the top of its documents.

    Years 1-4 lead with the bare title ("Computers Are Everywhere") above the
    running head; Years 5-7 print "Year 5 - Unit 3 - Networks and Digital
    Communication" or repeat "Unit 4 — Computer Systems" as the title line.
    """
    priority = {"Lesson": 0, "Reference": 1, "Activities": 2, "Practice": 3, "Other": 4}
    for document in sorted(documents, key=lambda item: priority.get(item["document_type"], 9)):
        texts = [block["text"] for block in document["blocks"][:12]]
        for text in texts:
            match = re.match(
                rf"^(?:Year\s+\d+\s*(?:Computing)?\s*[-–—]\s*)?Unit\s+(?:{unit}|\d+\.\d+)\s*[:\-–—]\s*(.+)$",
                text,
                re.IGNORECASE,
            )
            if match:
                candidate = clean(match.group(1))
                candidate = re.sub(r"^(?:Computing\s*[-–—]\s*)?", "", candidate)
                if candidate and not re.match(r"^(lesson|practice|activit|reference|teacher)", candidate, re.I):
                    return candidate
        for index, text in enumerate(texts[:3]):
            following = texts[index + 1] if index + 1 < len(texts) else ""
            if re.search(r"Year\s+\d+\s+Computing", following, re.I) and 2 < len(text) <= 90 \
                    and not re.match(r"^(year|unit|reference|practice|activities|lesson)", text, re.I):
                return clean(text)
        first = clean(texts[0]) if texts else ""
        if 2 < len(first) <= 90 and not re.match(
                r"^(year|unit|lesson|practice|activit|reference|teacher|how to use|name:)", first, re.I):
            return first
    return f"Unit {unit}"


def extract(downloads: Path) -> dict:
    grades: dict[int, dict] = {}
    archives = computing_archives(downloads)
    if not archives:
        raise SystemExit(f"No Computing archives found in {downloads}. Expected files named 'Year <n>-<stamp>-<part>.zip'.")
    for grade, zip_path in archives.items():
        documents = []
        with zipfile.ZipFile(zip_path) as archive:
            for entry in sorted(archive.infolist(), key=lambda item: item.filename):
                if not entry.filename.lower().endswith(".docx"):
                    continue
                unit_match = re.search(r"Unit\s+(\d+)", entry.filename, re.IGNORECASE)
                if not unit_match:
                    continue
                unit = int(unit_match.group(1))
                doc_type = document_type(entry.filename)
                documents.append(parse_document(archive.read(entry), entry.filename, grade, unit, doc_type))
        by_unit = defaultdict(list)
        for document in documents:
            by_unit[document["unit"]].append(document)

        # Years 1-4 ship a Teacher & Parent Guide, so their prose addresses the
        # adult and the builder has to rewrite it for the learner. Years 5-7
        # ship a student lesson book that can be carried across as written.
        audience = "teacher-guide" if any(
            "Teacher" in doc["source_file"] for doc in documents) else "student"

        units = []
        for unit_number in sorted(by_unit):
            unit_docs = by_unit[unit_number]
            counts = Counter(item["document_type"] for item in unit_docs)
            units.append({
                "grade": grade,
                "unit": unit_number,
                "unit_id": f"comp-g{grade:02d}-u{unit_number:02d}",
                "title": infer_unit_title(unit_docs, unit_number),
                "audience": audience,
                "lesson_documents": counts["Lesson"],
                "activity_documents": counts["Activities"],
                "practice_documents": counts["Practice"],
                "reference_documents": counts["Reference"],
                "other_documents": counts["Other"],
                "source_document_count": len(unit_docs),
                "source_block_count": sum(item["block_count"] for item in unit_docs),
                "review_status": "Imported - curriculum review required",
            })
        grades[grade] = {
            "metadata": {
                "schema_version": "Ehel Computing Content Package v1.0",
                "grade": grade,
                "grade_id": f"g{grade:02d}",
                "subject": "Computing",
                "audience": audience,
                "source_archive": zip_path.name,
                "unit_count": len(units),
                "source_document_count": len(documents),
            },
            "units": units,
            "documents": documents,
        }
    return {"grades": grades}


def sample(model: dict) -> None:
    for grade, package in model["grades"].items():
        print(f"\n### Grade {grade} ({package['metadata']['audience']})")
        seen = set()
        for document in package["documents"]:
            doc_type = document["document_type"]
            if doc_type in seen:
                continue
            seen.add(doc_type)
            print(f"\n[{doc_type}] {document['source_file']}")
            for block in document["blocks"][:18]:
                print(f"{block['sequence']:02d} {block['content_kind']} | {block['section'][:38]} | {block['text'][:150]}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--downloads", type=Path, default=Path.home() / "Downloads")
    parser.add_argument("--output", type=Path)
    parser.add_argument("--sample", action="store_true")
    args = parser.parse_args()
    model = extract(args.downloads)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"Wrote {args.output}")
    if args.sample:
        sample(model)
    print(json.dumps({grade: package["metadata"] for grade, package in model["grades"].items()}, indent=2))


if __name__ == "__main__":
    main()
