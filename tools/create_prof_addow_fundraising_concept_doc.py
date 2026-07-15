from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "outputs" / "template_model.docx"
OUT = ROOT / "outputs" / "Prof_Addow_TVET_Diaspora_Fundraising_Concept_Note.docx"


def clear_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def paragraph(text="", style=None, align=None, before=0, after=6, line=1.15):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


def heading(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=14, bold=True, color="2E74B5")
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Prof Addow TVET School | Diaspora Partnership for Skills, Jobs and Youth Opportunity")
    set_run_font(run, size=8.5, color="666666")


def add_bottom_rule(paragraph, color="2E74B5", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


doc = Document(str(TEMPLATE))
clear_body(doc)

section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
add_footer(section)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
for style_name in ("Heading 1", "List Bullet"):
    styles[style_name].font.name = "Calibri"

title = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2)
r = title.add_run("Prof Addow TVET School Diaspora Fundraising Initiative")
set_run_font(r, size=18, bold=True, color="1F4D78")

subtitle = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)
r = subtitle.add_run(
    "Concept Note: Mobilizing Somali Diaspora Support for Scholarships, Institutional Strengthening and Volunteer Engagement"
)
set_run_font(r, size=11.5, italic=True, color="444444")
add_bottom_rule(subtitle)

heading("Executive Summary")
paragraph(
    "Prof Addow TVET School in Galkayo plays a practical role in helping young people gain employable skills, build confidence and move toward productive livelihoods. Many students, however, face financial barriers that prevent them from enrolling, completing training or purchasing the tools and materials required for learning. This concept proposes a structured Somali diaspora fundraising and volunteer mobilization initiative to support student scholarships, direct institutional needs and professional volunteer engagement."
)
paragraph(
    "The initiative will create a trusted channel through which diaspora individuals, businesses, associations, mosques, alumni and professional networks can contribute once or twice a year through coordinated fundraising campaigns. It will also invite qualified diaspora volunteers to teach short modules, mentor students, support administration, assist with grant writing and connect graduates to employment or apprenticeship opportunities."
)

heading("Rationale")
paragraph(
    "TVET education is one of the most direct pathways from vulnerability to opportunity. In Galkayo, skills training can help youth, women and low-income learners access work in trades, services, entrepreneurship and community rebuilding. At the same time, the Somali diaspora has strong goodwill and significant professional capacity, but many supporters need a clear, accountable and practical mechanism for contributing to local education. A well-organized fundraising model can turn scattered goodwill into predictable support for the school and its students."
)

heading("Objectives")
bullet("Establish a transparent diaspora scholarship fund for vulnerable and high-potential students.")
bullet("Mobilize direct financial support for tools, equipment, learning materials, teacher development and school operations.")
bullet("Organize one major annual fundraising campaign and, where feasible, a second mid-year campaign to sustain predictable support.")
bullet("Create a diaspora volunteer programme for teaching, mentoring, administrative support, curriculum advice and career guidance.")
bullet("Strengthen trust through regular reporting, student stories, financial updates and visible results.")

heading("Strategic Activities")
bullet("Diaspora Fundraising Campaigns: Hold one flagship fundraising campaign each year, with an optional second campaign six months later during Ramadan, Eid season, summer diaspora gatherings or end-of-year giving.")
bullet("Student Scholarship Sponsorships: Offer clear sponsorship packages covering tuition, transport, tools, certification fees and basic learning materials.")
bullet("Direct School Support: Raise funds for workshop equipment, computers, internet access, classroom materials, solar power, facility improvements and teacher training.")
bullet("Volunteer Teaching and Mentoring: Recruit diaspora professionals to deliver short online or in-person sessions in technical skills, English, digital literacy, entrepreneurship, job readiness and life skills.")
bullet("Administrative and Governance Support: Engage volunteers to assist with finance systems, communications, fundraising records, policy development, proposal writing and monitoring.")
bullet("Diaspora Ambassadors Network: Identify trusted supporters in key diaspora cities to organize community events, business sponsorships and recurring donor commitments.")
bullet("Reporting and Accountability: Produce simple quarterly updates showing funds received, students supported, equipment purchased, volunteer hours delivered and graduate progress.")

heading("Implementation Framework")
paragraph(
    "The school will establish a small coordination team made up of school leadership, a diaspora advisory group, local community representatives and student-facing staff. This team will agree scholarship criteria, fundraising targets, campaign dates, reporting templates and volunteer roles. Donations should be tracked through a dedicated account or approved platform, with receipts, expenditure summaries and beneficiary updates shared with supporters."
)
paragraph(
    "The annual campaign will focus on a clear target, such as sponsoring a defined number of students or equipping a specific training department. The second fundraising round, if held, will focus on urgent student retention needs, tools and materials, or preparing the next training cohort."
)

heading("Expected Results")
bullet("Increased enrollment and retention of students from low-income and vulnerable households.")
bullet("More students completing market-relevant technical training with the tools needed to succeed.")
bullet("Improved workshops, learning materials, teacher capacity and administrative systems.")
bullet("Stronger links between students, diaspora professionals, employers and apprenticeship opportunities.")
bullet("A credible annual and semiannual fundraising rhythm that builds donor confidence and long-term sustainability.")

heading("Challenges")
paragraph(
    "Key risks include donor fatigue, limited trust in financial management, inconsistent volunteer availability, unclear scholarship selection, and weak follow-up after campaigns. These risks can be reduced through transparent reporting, defined campaign targets, published selection criteria, regular donor communication and a small but active diaspora advisory group."
)

heading("Opportunities")
paragraph(
    "The initiative can build on strong diaspora attachment to Somalia, the practical value of TVET education, community pride in Galkayo and the desire among professionals to give back. Even modest recurring contributions, when organized once or twice a year, can fund scholarships, improve workshops and create a visible pipeline from training to employment."
)

heading("Guiding Principles")
paragraph(
    "Somali ownership • Transparency • Inclusion • Student dignity • Gender and youth participation • Practical skills for employment • Volunteer service • Accountability to donors and the community."
)

heading("Conclusion")
paragraph(
    "The Prof Addow TVET School Diaspora Fundraising Initiative offers a practical, trusted and community-owned pathway to invest in young people in Galkayo. By combining scholarships, direct financial support and professional volunteer service, the Somali diaspora can help students access training, complete their studies and move toward dignified livelihoods. A disciplined once- or twice-yearly fundraising model will make support predictable, measurable and sustainable."
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(OUT)
