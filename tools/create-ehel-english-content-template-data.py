from __future__ import annotations

import argparse
import hashlib
import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "tmp" / "year2-curriculum" / "Year 2"
DEFAULT_VOCABULARY = ROOT / "src" / "prototypes" / "ehel-academy" / "vocabulary" / "grade2-vocabulary.json"
DEFAULT_OUTPUT = ROOT / "tmp" / "ehel-template-v1" / "grade2-content-model.json"


def clean(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def slug(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", value.lower().replace("'", "")).strip("-")


def vocabulary_starter(word: str, word_type: str, source_starter: str) -> str:
    if word_type in {"expression", "phrase"}:
        return f'I can say "{word}" when'
    if word_type == "number":
        return f"I can count to {word}"
    return source_starter


def multiline(values: list[str]) -> str:
    return "\n".join(value.strip() for value in values if value.strip())


def paragraphs(document: Document) -> list[dict[str, str]]:
    return [
        {"text": paragraph.text.strip(), "style": paragraph.style.name}
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]


def index_of(items: list[dict[str, str]], prefix: str, start: int = 0) -> int:
    for index in range(start, len(items)):
        if items[index]["text"].startswith(prefix):
            return index
    return len(items)


def section(items: list[dict[str, str]], start_prefix: str, end_prefix: str) -> list[dict[str, str]]:
    start = index_of(items, start_prefix)
    end = index_of(items, end_prefix, start + 1)
    return items[start + 1:end]


def split_heading_sections(items: list[dict[str, str]], pattern: str) -> list[tuple[str, list[dict[str, str]]]]:
    matcher = re.compile(pattern, re.IGNORECASE)
    headings = [index for index, item in enumerate(items) if matcher.match(item["text"])]
    output = []
    for position, start in enumerate(headings):
        end = headings[position + 1] if position + 1 < len(headings) else len(items)
        output.append((items[start]["text"], items[start + 1:end]))
    return output


def source_ref(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def unit_id(unit_number: int) -> str:
    return f"eng-g02-t{((unit_number - 1) // 3) + 1:02d}-u{unit_number:02d}"


def parse_story_metadata(document: Document) -> dict[str, str]:
    metadata = {"genre": "", "theme": "", "setting": ""}
    if not document.tables:
        return metadata
    text = document.tables[0].cell(0, 0).text
    for key in metadata:
        match = re.search(rf"{key}:\s*([^\n]+)", text, re.IGNORECASE)
        if match:
            metadata[key] = clean(match.group(1))
    return metadata


def parse_grammar_lesson(body: list[dict[str, str]]) -> dict[str, str]:
    labels = {
        "What it means and why it matters": "explanation",
        "How to use it": "rule_examples",
        "How to form it": "rule_examples",
        "The word pairs": "rule_examples",
        "A common mistake": "common_mistake",
        "Memory tip": "memory_tip",
    }
    fields: dict[str, list[str]] = {
        "explanation": [],
        "rule_examples": [],
        "common_mistake": [],
        "memory_tip": [],
        "practice": [],
    }
    current = "explanation"
    for item in body:
        text = item["text"]
        if text in labels:
            current = labels[text]
            continue
        if re.match(r"^Practice\s+\d+", text, re.IGNORECASE):
            current = "practice"
            continue
        fields[current].append(text)
    return {key: multiline(values) for key, values in fields.items()}


def extract_grade(source_dir: Path, vocabulary_path: Path) -> dict:
    vocabulary = json.loads(vocabulary_path.read_text(encoding="utf-8"))
    vocabulary_units = {unit["number"]: unit for unit in vocabulary["units"]}
    model = {
        "metadata": {
            "schema_version": "Ehel English Content Template v1.0",
            "grade_id": "g02",
            "grade": 2,
            "subject": "English",
            "term_mapping": "Provisional: Units 1-3 Term 1; Units 4-6 Term 2; Units 7-9 Term 3",
        },
        "units": [], "outcomes": [], "vocabulary": [], "readings": [], "comprehension": [],
        "grammar": [], "speaking": [], "writing": [], "activities": [], "assignments": [],
        "quizzes": [], "live_sessions": [], "teacher_notes": [], "answer_key": [],
        "rubrics": [], "self_assessment": [], "sources": [],
    }

    for unit_number in range(1, 10):
        uid = unit_id(unit_number)
        term_number = ((unit_number - 1) // 3) + 1
        unit_dir = source_dir / f"Unit {unit_number}"
        paths = {kind: unit_dir / f"Unit {unit_number} - {kind}.docx" for kind in ("Lesson", "Story", "Grammar", "Vocabulary")}
        documents = {kind: Document(path) for kind, path in paths.items()}
        lesson_items = paragraphs(documents["Lesson"])
        story_items = paragraphs(documents["Story"])
        grammar_items = paragraphs(documents["Grammar"])

        title = re.sub(r"^Unit\s+\d+:\s*", "", lesson_items[0]["text"])
        overview_index = index_of(lesson_items, "Unit Overview")
        overview = lesson_items[overview_index + 1]["text"] if overview_index + 1 < len(lesson_items) else ""
        path_items = section(lesson_items, "Your Learning Path", "Part 1:")
        model["units"].append({
            "grade_id": "g02", "subject": "English", "term_id": f"t{term_number:02d}",
            "unit_id": uid, "unit_number": unit_number, "unit_title": title,
            "unit_overview": overview, "learning_path": multiline([item["text"] for item in path_items]),
            "content_origin": "Authored source", "review_status": "Imported - curriculum review required",
            "source_file": source_ref(paths["Lesson"]),
        })

        outcome_start = index_of(lesson_items, "By the end of this unit")
        outcome_end = index_of(lesson_items, "Your Learning Path", outcome_start + 1)
        outcome_number = 0
        for item in lesson_items[outcome_start + 1:outcome_end]:
            if item["style"].startswith("List"):
                outcome_number += 1
                model["outcomes"].append({
                    "outcome_id": f"{uid}-lo{outcome_number:02d}", "unit_id": uid,
                    "sequence": outcome_number, "learning_outcome": item["text"],
                    "evidence_of_learning": "", "content_origin": "Authored source",
                    "review_status": "Imported - mapping required", "source_file": source_ref(paths["Lesson"]),
                })

        part1 = section(lesson_items, "Part 1:", "Part 2:")
        reading_sections = split_heading_sections(part1, r"^(Reading|Listening)(\s+\d+)?\s*(—|-)\s*")
        reading_number = 0
        for heading, body in reading_sections:
            reading_number += 1
            rid = f"{uid}-read{reading_number:02d}"
            kind = "Listening" if heading.lower().startswith("listening") else "Reading"
            title_text = re.split(r"\s+(?:—|-)\s+", heading, maxsplit=1)[-1]
            questions = [item["text"] for item in body if item["style"].startswith("List")]
            content = [item["text"] for item in body if not item["style"].startswith("List") and not item["text"].startswith("After reading")]
            model["readings"].append({
                "reading_id": rid, "unit_id": uid, "sequence": reading_number, "reading_type": kind,
                "title": title_text, "genre": "", "theme": "", "setting": "",
                "passage_or_script": multiline(content), "audio_required": kind == "Listening",
                "content_origin": "Authored source", "review_status": "Imported - editorial review required",
                "source_file": source_ref(paths["Lesson"]),
            })
            for question in questions:
                qnum = len([row for row in model["comprehension"] if row["unit_id"] == uid]) + 1
                model["comprehension"].append({
                    "question_id": f"{uid}-cq{qnum:03d}", "unit_id": uid, "reading_id": rid,
                    "section": kind, "sequence": qnum, "question_type": "Short answer",
                    "question": question, "correct_answer": "", "explanation": "", "marks": 1,
                    "outcome_id": "", "difficulty": "Core", "content_origin": "Authored source",
                    "review_status": "Answer required", "source_file": source_ref(paths["Lesson"]),
                })

        story_title = story_items[2]["text"] if len(story_items) > 2 else f"Unit {unit_number} story"
        story_after = index_of(story_items, "After Reading:")
        story_passage = [item["text"] for item in story_items[3:story_after]]
        story_meta = parse_story_metadata(documents["Story"])
        story_id = f"{uid}-story01"
        model["readings"].append({
            "reading_id": story_id, "unit_id": uid, "sequence": reading_number + 1, "reading_type": "Story",
            "title": story_title, "genre": story_meta["genre"], "theme": story_meta["theme"],
            "setting": story_meta["setting"], "passage_or_script": multiline(story_passage),
            "audio_required": True, "content_origin": "Authored source",
            "review_status": "Imported - editorial review required", "source_file": source_ref(paths["Story"]),
        })
        answer_start = index_of(story_items, "Answer these questions")
        think_start = index_of(story_items, "Think and talk", answer_start + 1)
        find_start = index_of(story_items, "Find the words", think_start + 1)
        writing_start = index_of(story_items, "Now You Try:", find_start + 1)
        story_question_groups = [
            ("Story comprehension", story_items[answer_start + 1:think_start]),
            ("Think and talk", story_items[think_start + 1:find_start]),
        ]
        for section_name, items in story_question_groups:
            for item in items:
                if not item["style"].startswith("List"):
                    continue
                qnum = len([row for row in model["comprehension"] if row["unit_id"] == uid]) + 1
                model["comprehension"].append({
                    "question_id": f"{uid}-cq{qnum:03d}", "unit_id": uid, "reading_id": story_id,
                    "section": section_name, "sequence": qnum, "question_type": "Short answer",
                    "question": item["text"], "correct_answer": "", "explanation": "", "marks": 1,
                    "outcome_id": "", "difficulty": "Core" if section_name == "Story comprehension" else "Stretch",
                    "content_origin": "Authored source", "review_status": "Answer required",
                    "source_file": source_ref(paths["Story"]),
                })

        speaking_items = section(lesson_items, "Part 2:", "Part 3:")
        for sequence, (heading, body) in enumerate(split_heading_sections(speaking_items, r"^Speaking"), 1):
            model["speaking"].append({
                "speaking_id": f"{uid}-speak{sequence:02d}", "unit_id": uid, "sequence": sequence,
                "activity_type": "Game" if "Game" in heading else "Speaking practice", "title": heading,
                "instructions_and_model_lines": multiline([item["text"] for item in body]),
                "recording_required": any("Record" in item["text"] for item in body),
                "ai_tutor_prompt": "", "outcome_id": "", "content_origin": "Authored source",
                "review_status": "Imported - mapping required", "source_file": source_ref(paths["Lesson"]),
            })

        writing_items = section(lesson_items, "Part 3:", "Part 4:")
        writing_sections = split_heading_sections(writing_items, r"^Writing\s+\d+")
        for sequence, (heading, body) in enumerate(writing_sections, 1):
            model["writing"].append({
                "writing_id": f"{uid}-write{sequence:02d}", "unit_id": uid, "sequence": sequence,
                "title": heading, "prompt_and_instructions": multiline([item["text"] for item in body]),
                "model_text": "", "sentence_starter": "", "expected_length": "Teacher to confirm",
                "rubric_id": "rub-writing-v1", "outcome_id": "", "content_origin": "Authored source",
                "review_status": "Imported - rubric mapping required", "source_file": source_ref(paths["Lesson"]),
            })
        story_writing = [item["text"] for item in story_items[writing_start + 1:] if item["style"].startswith("List")]
        if story_writing:
            sequence = len(writing_sections) + 1
            model["writing"].append({
                "writing_id": f"{uid}-write{sequence:02d}", "unit_id": uid, "sequence": sequence,
                "title": f"Story response: {story_title}", "prompt_and_instructions": multiline(story_writing),
                "model_text": "", "sentence_starter": "", "expected_length": "Teacher to confirm",
                "rubric_id": "rub-writing-v1", "outcome_id": "", "content_origin": "Authored source",
                "review_status": "Imported - rubric mapping required", "source_file": source_ref(paths["Story"]),
            })

        activity_items = section(lesson_items, "Part 4:", "Self-Study Tips")
        for sequence, (heading, body) in enumerate(split_heading_sections(activity_items, r"^Activity\s+\d+"), 1):
            answer_lines = [item["text"] for item in body if item["text"].startswith("Answers:")]
            instruction_lines = [item["text"] for item in body if not item["text"].startswith("Answers:")]
            activity_id = f"{uid}-act{sequence:02d}"
            model["activities"].append({
                "activity_id": activity_id, "unit_id": uid, "sequence": sequence, "title": heading,
                "activity_type": "Independent practice", "instructions_and_items": multiline(instruction_lines),
                "answer_summary": multiline(answer_lines).removeprefix("Answers: "), "outcome_id": "",
                "delivery_mode": "Online or workbook", "content_origin": "Authored source",
                "review_status": "Imported - interaction design required", "source_file": source_ref(paths["Lesson"]),
            })
            if answer_lines:
                model["answer_key"].append({
                    "answer_id": f"{activity_id}-answer", "unit_id": uid, "content_id": activity_id,
                    "content_type": "Activity", "answer_or_guidance": multiline(answer_lines).removeprefix("Answers: "),
                    "content_origin": "Authored source", "review_status": "Imported",
                    "source_file": source_ref(paths["Lesson"]),
                })

        grammar_core = grammar_items[:index_of(grammar_items, "Answer Key")]
        grammar_sections = split_heading_sections(grammar_core, r"^Lesson\s+\d+:")
        for sequence, (heading, body) in enumerate(grammar_sections, 1):
            parsed = parse_grammar_lesson(body)
            model["grammar"].append({
                "grammar_id": f"{uid}-grammar{sequence:02d}", "unit_id": uid, "sequence": sequence,
                "title": re.sub(r"^Lesson\s+\d+:\s*", "", heading), **parsed,
                "outcome_id": "", "content_origin": "Authored source",
                "review_status": "Imported - interaction design required", "source_file": source_ref(paths["Grammar"]),
            })
        answer_index = index_of(grammar_items, "Answer Key")
        answer_body = grammar_items[answer_index + 1:]
        answer_sections = split_heading_sections(answer_body, r"^Practice\s+\d+")
        for sequence, (heading, body) in enumerate(answer_sections, 1):
            model["answer_key"].append({
                "answer_id": f"{uid}-grammar{sequence:02d}-answer", "unit_id": uid,
                "content_id": f"{uid}-grammar{sequence:02d}", "content_type": "Grammar practice",
                "answer_or_guidance": multiline([item["text"] for item in body]),
                "content_origin": "Authored source", "review_status": "Imported",
                "source_file": source_ref(paths["Grammar"]),
            })

        vocab_unit = vocabulary_units[unit_number]
        for group_number, group in enumerate(vocab_unit["groups"], 1):
            for sequence, word in enumerate(group["words"], 1):
                model["vocabulary"].append({
                    "vocabulary_id": f"{uid}-v{group_number:02d}-{sequence:03d}-{slug(word['word'])}",
                    "unit_id": uid, "group_id": f"{uid}-vg{group_number:02d}",
                    "group_title": group["title"], "sequence": sequence, "word": word["word"],
                    "word_type": word["type"], "source_type": word["sourceType"],
                    "simple_meaning": word["meaning"], "example_sentence": word["example"],
                    "pronunciation_text": word["word"], "picture_needed": word["type"] in {"noun", "verb", "adjective", "position"},
                    "student_sentence_starter": vocabulary_starter(word["word"], word["type"], word["starter"]),
                    "spelling_practice": word["word"],
                    "ai_tutor_prompt": word["tutorPrompt"],
                    **{f"sentence_{index + 1}": sentence for index, sentence in enumerate(word["sentences"])},
                    "content_origin": "Authored core + generated enrichment",
                    "review_status": "Generated enrichment review required",
                    "source_file": source_ref(paths["Vocabulary"]),
                })

        self_start = index_of(lesson_items, "Self-Assessment:")
        self_end = min(index_of(lesson_items, "Getting Ready", self_start + 1), index_of(lesson_items, "Congratulations", self_start + 1))
        self_sequence = 0
        for item in lesson_items[self_start + 1:self_end]:
            if not item["style"].startswith("List"):
                continue
            self_sequence += 1
            statement = re.sub(r"\s+Yes\s+□\s+Not yet\s+□.*$", "", item["text"])
            model["self_assessment"].append({
                "self_assessment_id": f"{uid}-self{self_sequence:02d}", "unit_id": uid,
                "sequence": self_sequence, "statement": statement, "scale": "Not yet | With help | By myself",
                "outcome_id": f"{uid}-lo{min(self_sequence, max(outcome_number, 1)):02d}",
                "content_origin": "Authored source - scale normalized", "review_status": "Imported",
                "source_file": source_ref(paths["Lesson"]),
            })

        tutor_notes = []
        for kind, document in documents.items():
            for table in document.tables:
                text = clean(table.cell(0, 0).text)
                if "Ask Your AI Tutor" in text:
                    tutor_notes.append(f"{kind}: {text}")
        model["teacher_notes"].append({
            "teacher_note_id": f"{uid}-note01", "unit_id": uid, "note_type": "AI tutor prompts from source",
            "note": multiline(tutor_notes), "visibility": "Teacher", "content_origin": "Authored source",
            "review_status": "Imported - teacher adaptation required", "source_file": "Multiple unit files",
        })
        model["teacher_notes"].append({
            "teacher_note_id": f"{uid}-note02", "unit_id": uid, "note_type": "Content gap",
            "note": "Confirm term placement; map every activity and assessment to outcomes; approve comprehension answers; approve generated quiz, assignment, live-session and rubric mappings.",
            "visibility": "Curriculum team", "content_origin": "Migration note", "review_status": "Action required",
            "source_file": "Ehel English Content Template v1 migration",
        })

        model["assignments"].append({
            "assignment_id": f"{uid}-assignment01", "unit_id": uid, "title": f"{title} learning portfolio",
            "instructions": "Submit one completed writing task and one speaking recording from this unit.",
            "submission_type": "Writing + audio", "marks": 20, "outcome_ids": "Teacher to map",
            "rubric_ids": "rub-writing-v1 | rub-speaking-v1", "content_origin": "Generated scaffold",
            "review_status": "Teacher approval required", "source_file": "Derived from unit writing and speaking tasks",
        })

        unit_words = [row for row in model["vocabulary"] if row["unit_id"] == uid]
        for sequence, word in enumerate(unit_words[:5], 1):
            other_meanings = [item["simple_meaning"] for item in unit_words if item["vocabulary_id"] != word["vocabulary_id"]]
            options = [word["simple_meaning"], *other_meanings[:3]]
            shift = (sequence - 1) % len(options)
            options = options[shift:] + options[:shift]
            model["quizzes"].append({
                "quiz_id": f"{uid}-quiz01", "question_id": f"{uid}-quiz01-q{sequence:02d}",
                "unit_id": uid, "quiz_title": f"{title} vocabulary checkpoint", "sequence": sequence,
                "question_type": "Multiple choice", "question": f"What does '{word['word']}' mean?",
                "options": " | ".join(options), "correct_answer": word["simple_meaning"],
                "explanation": word["example_sentence"], "marks": 1, "outcome_id": "Teacher to map",
                "difficulty": "Core", "content_origin": "Generated draft from approved vocabulary",
                "review_status": "Teacher approval required", "source_file": source_ref(paths["Vocabulary"]),
            })

        for session_number, focus in enumerate(("Vocabulary, reading and speaking workshop", "Writing, grammar and feedback clinic"), 1):
            model["live_sessions"].append({
                "live_session_id": f"{uid}-live{session_number:02d}", "unit_id": uid,
                "session_number": session_number, "title": focus, "duration_minutes": 45,
                "before_session": "Complete the related self-paced lesson and bring one question.",
                "agenda": "5 min welcome; 10 min retrieval; 20 min guided practice; 5 min learner demonstration; 5 min next steps.",
                "after_session": "Apply teacher feedback and update the unit portfolio.",
                "outcome_ids": "Teacher to map", "content_origin": "Generated scaffold",
                "review_status": "Teacher approval required", "source_file": "Derived from source recommendation: twice-weekly teacher sessions",
            })

        for kind, path in paths.items():
            model["sources"].append({
                "source_id": f"{uid}-{kind.lower()}", "unit_id": uid, "source_type": kind,
                "source_file": source_ref(path), "size_bytes": path.stat().st_size,
                "sha256": hashlib.sha256(path.read_bytes()).hexdigest(), "import_status": "Imported",
            })

    rubric_definitions = {
        "rub-speaking-v1": ("Speaking", [
            ("Clarity", "Hard to hear or understand", "Sometimes clear", "Mostly clear", "Clear and confident"),
            ("Vocabulary", "Uses very few target words", "Uses one target word", "Uses several target words", "Uses target words accurately and naturally"),
            ("Sentence accuracy", "Needs full support", "Uses partial sentences", "Uses mostly complete sentences", "Uses complete accurate sentences"),
            ("Interaction", "Does not respond yet", "Responds with support", "Responds and takes a turn", "Responds, asks and extends ideas"),
        ]),
        "rub-writing-v1": ("Writing", [
            ("Ideas", "Ideas are not yet clear", "One relevant idea", "Several relevant ideas", "Clear, developed and engaging ideas"),
            ("Organization", "Needs support to order ideas", "Some order is visible", "Ideas follow a clear order", "Ideas connect smoothly"),
            ("Language", "Very limited target language", "Some target words and patterns", "Mostly accurate target language", "Accurate and varied target language"),
            ("Conventions", "Capital letters and punctuation need support", "Some correct conventions", "Mostly correct conventions", "Consistently correct conventions"),
        ]),
        "rub-reading-v1": ("Reading fluency", [
            ("Accuracy", "Many words need support", "Some words need support", "Most words are accurate", "Accurate throughout"),
            ("Pace", "Reads word by word", "Uneven pace", "Mostly smooth pace", "Smooth and natural pace"),
            ("Expression", "Little expression yet", "Some expression", "Expression supports meaning", "Confident expressive reading"),
            ("Understanding", "Needs support to retell", "Retells one detail", "Retells main ideas", "Retells and explains important details"),
        ]),
    }
    for rubric_id, (target, criteria) in rubric_definitions.items():
        for sequence, criterion in enumerate(criteria, 1):
            model["rubrics"].append({
                "rubric_id": rubric_id, "target": target, "criterion_id": f"{rubric_id}-c{sequence:02d}",
                "criterion": criterion[0], "level_1": criterion[1], "level_2": criterion[2],
                "level_3": criterion[3], "level_4": criterion[4], "maximum_marks": 4,
                "content_origin": "Ehel reusable template", "review_status": "Curriculum approval required",
            })
    return model


def main() -> None:
    parser = argparse.ArgumentParser(description="Build normalized Grade 2 data for Ehel English Content Template v1.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--vocabulary", type=Path, default=DEFAULT_VOCABULARY)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    model = extract_grade(args.source, args.vocabulary)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {args.output}")
    for key, value in model.items():
        if isinstance(value, list):
            print(f"{key}: {len(value)}")


if __name__ == "__main__":
    main()
